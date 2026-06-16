import docx

def main():
    path = r'assets/products material descirption.docx'
    doc = docx.Document(path)
    print("Document Paragraphs count:", len(doc.paragraphs))
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip():
            print(f"Para {idx}: {p.text}")
            
    # Also print tables if any
    print("Tables count:", len(doc.tables))
    for t_idx, table in enumerate(doc.tables):
        print(f"--- Table {t_idx} ---")
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            print(row_data)

if __name__ == '__main__':
    main()
