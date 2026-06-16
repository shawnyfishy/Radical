import docx

def main():
    path = r'assets/products material descirption.docx'
    doc = docx.Document(path)
    
    print("Searching in document paragraphs:")
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            for keyword in ["monument", "monarch", "eclipse", "duality"]:
                if keyword in text.lower():
                    print(f"Para {idx}: {text}")
                    
    print("\nSearching in document tables:")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            for keyword in ["monument", "monarch", "eclipse", "duality"]:
                if keyword in row_text.lower():
                    print(f"Table {t_idx}, Row {r_idx}: {row_text}")

if __name__ == '__main__':
    main()
